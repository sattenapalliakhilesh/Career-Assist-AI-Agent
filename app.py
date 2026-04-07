import os
import json
import smtplib
import threading
import time
import requests
from flask import Flask, render_template, request, jsonify
from anthropic import Anthropic
from dotenv import load_dotenv
from bs4 import BeautifulSoup
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from datetime import datetime
from apscheduler.schedulers.background import BackgroundScheduler
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

load_dotenv()

app = Flask(__name__)
client = Anthropic()

SERPER_API_KEY = os.getenv("SERPER_API_KEY")
SCHEDULER_ENABLED = os.getenv("SCHEDULER_ENABLED", "false").lower() == "true"

def load_resume():
    with open("resume.txt", "r") as f:
        return f.read()

def search_web(query):
    url = "https://google.serper.dev/search"
    headers = {"X-API-KEY": SERPER_API_KEY, "Content-Type": "application/json"}
    payload = {"q": query, "num": 3, "tbs": "qdr:w"}
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=10)
        results = response.json()
        output = []
        for r in results.get("organic", [])[:3]:
            output.append(f"Title: {r.get('title')}\nSnippet: {r.get('snippet')}\nURL: {r.get('link')}\nDate: {r.get('date', 'Unknown')}")
        return "\n\n".join(output)
    except Exception as e:
        return f"Search failed: {str(e)}"

def fetch_url(url):
    try:
        response = requests.get(url, timeout=8, headers={"User-Agent": "Mozilla/5.0"})
        soup = BeautifulSoup(response.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        return text[:2000]
    except Exception as e:
        return f"Could not fetch: {str(e)}"

TOOLS = [
    {
        "name": "search_web",
        "description": "Search LinkedIn for job postings in Canada only. Results filtered to last 7 days.",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {"type": "string"}
            },
            "required": ["query"]
        }
    },
    {
        "name": "fetch_url",
        "description": "Fetch a LinkedIn job posting page.",
        "input_schema": {
            "type": "object",
            "properties": {
                "url": {"type": "string"}
            },
            "required": ["url"]
        }
    }
]

def run_tool(tool_name, tool_input):
    if tool_name == "search_web":
        return search_web(tool_input["query"])
    elif tool_name == "fetch_url":
        url = tool_input["url"]
        if "linkedin.com" not in url:
            return "Blocked — only LinkedIn URLs are allowed."
        return fetch_url(url)
    return "Tool not found"

def call_claude(model, system_prompt, user_message, max_tokens, use_tools=False):
    messages = [{"role": "user", "content": user_message}]
    agent_log = []
    max_iterations = 5
    iteration = 0

    while iteration < max_iterations:
        iteration += 1
        kwargs = {
            "model": model,
            "max_tokens": max_tokens,
            "system": system_prompt,
            "messages": messages
        }
        if use_tools:
            kwargs["tools"] = TOOLS

        for attempt in range(3):
            try:
                response = client.messages.create(**kwargs)
                break
            except Exception as e:
                if "429" in str(e) and attempt < 2:
                    wait = (attempt + 1) * 30
                    print(f"Rate limit hit — waiting {wait}s before retry")
                    time.sleep(wait)
                else:
                    raise e

        if response.stop_reason == "tool_use":
            tool_results = []
            for block in response.content:
                if block.type == "tool_use":
                    log_entry = f"Searching: {list(block.input.values())[0][:80]}"
                    agent_log.append(log_entry)
                    print(log_entry)
                    result = run_tool(block.name, block.input)
                    tool_results.append({
                        "type": "tool_result",
                        "tool_use_id": block.id,
                        "content": result[:2000]
                    })
            messages.append({"role": "assistant", "content": response.content})
            messages.append({"role": "user", "content": tool_results})

        elif response.stop_reason == "end_turn":
            final_text = ""
            for block in response.content:
                if hasattr(block, "text"):
                    final_text += block.text
            return final_text, agent_log

        else:
            break

    return "", agent_log

def parse_json(text):
    clean = text.strip()
    if "```" in clean:
        parts = clean.split("```")
        for part in parts:
            p = part.strip()
            if p.startswith("json"):
                p = p[4:].strip()
            if p.startswith("{"):
                clean = p
                break
    start = clean.find("{")
    end = clean.rfind("}") + 1
    if start >= 0 and end > start:
        clean = clean[start:end]
    return json.loads(clean)

def _salary_meets_minimum(salary_str):
    """Return True if the salary string indicates >= $75/hr or >= $180,000/year.
    Returns True (pass-through) when the salary cannot be parsed."""
    import re
    s = salary_str.lower().replace(",", "").replace("$", "").replace("cad", "").strip()

    # Extract all numbers from the string
    numbers = [float(n) for n in re.findall(r"\d+(?:\.\d+)?", s)]
    if not numbers:
        return True  # unparseable — let it through

    # Determine if hourly or annual
    is_hourly = any(word in s for word in ["hour", "hr", "/h", "per h"])
    is_annual = any(word in s for word in ["year", "yr", "annual", "salary", "k"])

    # Handle shorthand like "180k" or "200k"
    if "k" in s:
        numbers = [n * 1000 if n < 10000 else n for n in numbers]

    max_val = max(numbers)

    if is_hourly:
        return max_val >= 75
    if is_annual:
        return max_val >= 180000

    # Ambiguous — use magnitude to guess
    if max_val >= 1000:
        return max_val >= 180000  # treat as annual
    return max_val >= 75  # treat as hourly


def find_jobs_only(criteria):
    system_prompt = """You are a job search specialist searching LinkedIn for jobs in Canada.

Search strategy — run these searches in this order:
1. site:linkedin.com/jobs "Scrum Master" Canada
2. site:linkedin.com/jobs "Project Manager" Canada
3. site:linkedin.com/jobs "Scrum Master" remote Canada
4. site:linkedin.com/jobs "Project Manager" remote Canada

Important rules:
- Only include results from linkedin.com
- Only include Scrum Master or Project Manager roles
- Only include jobs in Canada
- Only include jobs that are ACTIVE and currently accepting applications — skip closed, expired, or filled postings
- If salary is not mentioned that is okay — still include the job
- If salary IS mentioned, only include jobs paying at least $75/hour or $180,000 annually

PRIORITIZATION — sort jobs by how recently they were posted:
- Today or posted X hours ago = highest priority
- Yesterday or 1 day ago = high priority
- 2-3 days ago = medium priority
- 4-7 days ago = lower priority
- Over 7 days ago = exclude enittirely

Always return jobs sorted from newest to oldest.
Include the posting age clearly in posted_date field e.g. "2 hours ago", "1 day ago", "3 days ago".

Return whatever you find — even if only 1 or 2 jobs. Return ONLY this JSON:
{
    "jobs": [
        {
            "title": "Job title",
            "company": "Company name",
            "location": "City Province or Remote Canada",
            "url": "linkedin.com URL",
            "salary": "Salary if mentioned or empty string",
            "description": "One sentence about the role",
            "posted_date": "e.g. 2 hours ago / 1 day ago / 3 days ago",
            "posted_priority": 1,
            "is_active": true,
            "source": "linkedin"
        }
    ],
    "search_summary": "One sentence — what was searched, how many found, newest posting age"
}

posted_priority scoring:
- Today / hours ago = 1 (highest)
- 1 day ago = 2
- 2 days ago = 3
- 3 days ago = 4
- 4-7 days ago = 5
- Unknown date = 6 (lowest)

Sort the jobs array by posted_priority ascending — lowest number first.

If absolutely no jobs are found return:
{
    "jobs": [],
    "search_summary": "No matching jobs found on LinkedIn Canada this run"
}"""

    text, agent_log = call_claude(
        model="claude-haiku-4-5",
        system_prompt=system_prompt,
        user_message=f"Find Scrum Master and Project Manager jobs on LinkedIn in Canada, newest first:\n\n{criteria}",
        max_tokens=1500,
        use_tools=True
    )

    print(f"Raw agent response: {text[:300]}")

    if not text or not text.strip():
        return {
            "jobs": [],
            "search_summary": "Agent returned no response — will retry next run",
            "agent_log": agent_log
        }

    try:
        result = parse_json(text)

        filtered_jobs = []
        allowed_titles = ["scrum master", "project manager"]
        for job in result.get("jobs", []):
            url = job.get("url", "").lower()
            title = job.get("title", "").lower()
            location = job.get("location", "").lower()

            if "linkedin.com" not in url:
                print(f"Filtered — not LinkedIn: {job.get('title')} — {url}")
                continue
            if "united states" in location or ", us" in location or "usa" in location:
                print(f"Filtered — not Canada: {job.get('title')} — {location}")
                continue
            if not any(t in title for t in allowed_titles):
                print(f"Filtered — wrong role: {job.get('title')}")
                continue

            if not job.get("is_active", True):
                print(f"Filtered — not active/accepting applications: {job.get('title')} — {job.get('company')}")
                continue

            salary_raw = job.get("salary", "").strip()
            if salary_raw and not _salary_meets_minimum(salary_raw):
                print(f"Filtered — salary below minimum: {job.get('title')} — {salary_raw}")
                continue

            filtered_jobs.append(job)

        filtered_jobs.sort(key=lambda x: x.get("posted_priority", 6))

        result["jobs"] = filtered_jobs
        result["agent_log"] = agent_log

        if not filtered_jobs:
            result["search_summary"] = "Jobs were found but filtered out — try again later or broaden criteria"

        return result

    except Exception as e:
        print(f"Parse error in find_jobs_only: {e}\nRaw: {text[:500]}")
        return {
            "jobs": [],
            "search_summary": "Could not parse results — will retry next run",
            "agent_log": agent_log,
            "error": str(e)
        }

def tailor_for_job(job, resume):
    resume_short = resume[:3000]

    system_prompt = """You are an expert resume writer. Tailor the resume for this specific job.
Be concise. Return ONLY valid JSON — no other text:
{
    "match_score": 85,
    "match_summary": "One sentence why this is a good match",
    "key_requirements": ["req 1", "req 2", "req 3"],
    "missing_skills": ["skill 1"],
    "tailored_resume": "Full tailored resume — see format rules below",
    "cover_letter": "Short focused cover letter",
    "keywords_added": ["keyword 1", "keyword 2"]
}

tailored_resume MUST follow this exact plain-text structure:
Line 1: Candidate full name (e.g. AKHILESH SATTENAPALLI)
Line 2: Title | Certifications (e.g. Technical Program Manager | PMP | Harvard Business School Certified)
Line 3: Location | email | LinkedIn | phone
(blank line)
Then sections using these rules:
- Section headers in ALL CAPS on their own line (e.g. EXECUTIVE PROFILE, SELECTED CAREER HIGHLIGHTS, PROFESSIONAL LEADERSHIP HISTORY, CERTIFICATIONS, EDUCATION)
- Job entries on ONE line: Title | Company | City  Month Year – Month Year  (date at end separated by two spaces)
- Bullet points starting with • (one bullet per line)
- Use **bold text** around key phrases in paragraph sections
- Keep original sections and ordering from the source resume"""

    user_message = f"""Job: {job.get('title')} at {job.get('company')}
Location: {job.get('location')}
Description: {job.get('description')}
Salary: {job.get('salary', 'Not specified')}
Posted: {job.get('posted_date', 'Within last 7 days')}

Resume:
{resume_short}"""

    text, _ = call_claude(
        model="claude-sonnet-4-6",
        system_prompt=system_prompt,
        user_message=user_message,
        max_tokens=2500,
        use_tools=False
    )

    try:
        return parse_json(text)
    except Exception as e:
        print(f"Parse error in tailor_for_job: {e}\nRaw: {text[:300]}")
        return {
            "match_score": 0,
            "match_summary": "Could not tailor resume — please try again",
            "key_requirements": [],
            "missing_skills": [],
            "tailored_resume": "",
            "cover_letter": "",
            "keywords_added": [],
            "error": str(e)
        }

def find_and_tailor_jobs(criteria):
    resume = load_resume()
    search_result = find_jobs_only(criteria)

    if "error" in search_result and "jobs" not in search_result:
        return search_result

    jobs = search_result.get("jobs", [])
    agent_log = search_result.get("agent_log", [])
    tailored_jobs = []

    for i, job in enumerate(jobs):
        print(f"Tailoring {i+1}/{len(jobs)}: {job.get('title')} at {job.get('company')}")
        agent_log.append(f"Tailoring resume for: {job.get('title')} at {job.get('company')}")
        tailored = tailor_for_job(job, resume)
        merged = {**job, **tailored}
        tailored_jobs.append(merged)
        if i < len(jobs) - 1:
            time.sleep(15)

    return {
        "jobs": tailored_jobs,
        "search_summary": search_result.get("search_summary", ""),
        "agent_log": agent_log
    }

def analyze_job(job_input):
    resume = load_resume()
    resume_short = resume[:3000]

    system_prompt = f"""You are an expert resume writer. Analyze this job and tailor the resume.

Resume:
{resume_short}

Return ONLY valid JSON — no other text:
{{
    "match_score": 85,
    "match_summary": "2 sentence explanation",
    "key_requirements": ["req 1", "req 2", "req 3"],
    "missing_skills": ["skill 1"],
    "tailored_resume": "Full tailored resume",
    "cover_letter": "Full cover letter",
    "keywords_added": ["keyword 1", "keyword 2"]
}}"""

    text, agent_log = call_claude(
        model="claude-sonnet-4-6",
        system_prompt=system_prompt,
        user_message=f"Analyze and tailor my resume for:\n\n{job_input[:3000]}",
        max_tokens=2500,
        use_tools=True
    )

    try:
        result = parse_json(text)
        result["agent_log"] = agent_log
        return result
    except Exception as e:
        print(f"Parse error in analyze_job: {e}\nRaw: {text[:300]}")
        return {
            "error": f"Could not parse response. Please try again. ({str(e)})",
            "agent_log": agent_log
        }

def create_resume_doc(job, resume_text):
    import re

    doc = Document()

    # Page margins
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)

    # Base font
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    def _hr():
        """Thin horizontal rule as a paragraph bottom border."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _section_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)

    def _job_header(title_part, date_str):
        """Bold title left-aligned, italic date right-aligned via tab stop."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        r1 = p.add_run(title_part)
        r1.bold = True
        r1.font.name = 'Calibri'
        r1.font.size = Pt(11)
        if date_str:
            p.add_run('\t')
            r2 = p.add_run(date_str)
            r2.italic = True
            r2.font.name = 'Calibri'
            r2.font.size = Pt(11)

    def _bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

    def _body_paragraph(text):
        """Justified paragraph with **bold** inline support."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        parts = re.split(r'\*\*(.*?)\*\*', text)
        for idx, part in enumerate(parts):
            run = p.add_run(part)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            if idx % 2 == 1:
                run.bold = True

    # Regex: detects a date range at the end of a job entry line
    DATE_RE = re.compile(
        r'((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|'
        r'Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}'
        r'\s*[–\-]\s*'
        r'(?:Present|\d{4}|(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|'
        r'Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|'
        r'Dec(?:ember)?)\s+\d{4}))'
    )

    # ── Header block: first 3 non-empty lines = name / title+creds / contact ──
    lines = [ln.strip() for ln in resume_text.split('\n')]
    header_lines = []
    body_start = 0
    for idx, ln in enumerate(lines):
        if ln:
            header_lines.append(ln)
            if len(header_lines) == 3:
                body_start = idx + 1
                break

    # Name — large bold
    if len(header_lines) > 0:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(header_lines[0])
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(18)

    # Title / credentials — bold
    if len(header_lines) > 1:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(header_lines[1])
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

    # Contact line — normal
    if len(header_lines) > 2:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(header_lines[2])
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

    _hr()

    # ── Body ──
    EXPERIENCE_KEYWORDS = {'HISTORY', 'EXPERIENCE', 'EMPLOYMENT', 'CAREER'}
    first_section = True
    in_experience = False
    first_job_in_section = True

    i = body_start
    while i < len(lines):
        line = lines[i]
        i += 1

        if not line:
            continue

        # Section header: short ALL-CAPS line with at least one letter, no bullet prefix
        is_section = (
            line.upper() == line
            and len(line) < 65
            and any(c.isalpha() for c in line)
            and not line[0] in ('•', '-', '*')
        )
        if is_section:
            if not first_section:
                _hr()
            _section_header(line)
            first_section = False
            in_experience = bool(EXPERIENCE_KEYWORDS & set(line.upper().split()))
            first_job_in_section = True
            continue

        # Bullet point
        if line[0] in ('•', '-', '*'):
            _bullet(line.lstrip('•-* '))
            continue

        # Job entry header (contains a date range)
        m = DATE_RE.search(line)
        if m:
            if in_experience and not first_job_in_section:
                _hr()
            date_str = m.group(1)
            title_part = line[:m.start()].strip().rstrip('–- ').strip()
            _job_header(title_part, date_str)
            first_job_in_section = False
            continue

        # Regular body paragraph
        _body_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def create_cover_letter_doc(job, cover_text):
    import re

    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    # Date
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.add_run(datetime.now().strftime('%B %d, %Y')).font.size = Pt(11)

    # RE line
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"Re: {job.get('title')} — {job.get('company')}")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    # Divider
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Body
    for line in cover_text.split('\n'):
        line = line.strip()
        if not line:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(8)
        parts = re.split(r'\*\*(.*?)\*\*', line)
        for idx, part in enumerate(parts):
            run = p.add_run(part)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            if idx % 2 == 1:
                run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def send_email(subject, html_body, attachments=[]):
    gmail = os.getenv("GMAIL_ADDRESS")
    password = os.getenv("GMAIL_APP_PASSWORD")
    recipient = os.getenv("NOTIFY_EMAIL")
    msg = MIMEMultipart("mixed")
    msg["From"] = gmail
    msg["To"] = recipient
    msg["Subject"] = subject
    msg.attach(MIMEText(html_body, "html"))
    for filename, content in attachments:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(content)
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f"attachment; filename={filename}")
        msg.attach(part)
    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(gmail, password)
        server.sendmail(gmail, recipient, msg.as_string())
    print(f"Email sent to {recipient}")

def build_email_html(jobs, search_summary):
    html = f"""<html><body style="font-family:Arial,sans-serif;max-width:800px;margin:0 auto;padding:20px;">
<h1 style="color:#3b5bdb;">Daily Job Report — LinkedIn Canada</h1>
<p style="color:#666;">{datetime.now().strftime('%A, %B %d %Y')} — {len(jobs)} jobs found, sorted newest first</p>
<p style="background:#f8fafc;padding:12px;border-radius:8px;color:#555;font-size:14px;">{search_summary}</p>
<p style="font-size:13px;color:#888;">Roles: Scrum Master, Project Manager | Canada only | Posted last 7 days | Newest first</p>
<hr style="border:none;border-top:1px solid #e2e8f0;margin:20px 0;">"""

    for i, job in enumerate(jobs):
        score = job.get("match_score", 0)
        priority = job.get("posted_priority", 6)
        color = "#16a34a" if score >= 70 else "#d97706" if score >= 50 else "#dc2626"
        bg = "#f0fdf4" if score >= 70 else "#fffbeb" if score >= 50 else "#fff5f5"
        age_color = "#16a34a" if priority <= 1 else "#d97706" if priority <= 3 else "#888"
        company_clean = job.get('company', '').replace(' ', '_')

        freshness = ""
        if priority == 1:
            freshness = "🟢 Posted today"
        elif priority == 2:
            freshness = "🟡 Posted yesterday"
        elif priority <= 4:
            freshness = f"🟠 Posted {job.get('posted_date', 'recently')}"
        else:
            freshness = f"⚪ Posted {job.get('posted_date', 'recently')}"

        html += f"""
<div style="border:1px solid #e2e8f0;border-radius:12px;margin-bottom:20px;overflow:hidden;">
  <div style="padding:16px 20px;background:#f8fafc;border-bottom:1px solid #e2e8f0;">
    <div style="display:flex;justify-content:space-between;align-items:flex-start;">
      <div>
        <h2 style="font-size:16px;color:#111;margin:0 0 4px;">#{i+1} {job.get('title', '')}</h2>
        <p style="font-size:13px;color:#666;margin:0;">{job.get('company', '')} · {job.get('location', '')} {('· ' + job.get('salary')) if job.get('salary') else ''}</p>
      </div>
      <span style="background:{bg};color:{color};font-weight:700;padding:4px 12px;border-radius:20px;font-size:12px;white-space:nowrap;">{score}% match</span>
    </div>
    <p style="font-size:12px;color:{age_color};margin:8px 0 0;font-weight:600;">{freshness}</p>
  </div>
  <div style="padding:16px 20px;">
    <p style="font-size:13px;color:#555;margin-bottom:10px;">{job.get('match_summary', '')}</p>
    {f'<p><a href="{job.get("url")}" style="background:#3b5bdb;color:#fff;padding:8px 16px;border-radius:7px;text-decoration:none;font-size:13px;">View on LinkedIn</a></p>' if job.get('url') else ''}
    <p style="font-size:12px;color:#888;margin-top:8px;">Attachments: Job_{i+1}_{company_clean}_Resume.docx and Job_{i+1}_{company_clean}_CoverLetter.docx</p>
  </div>
</div>"""

    html += "</body></html>"
    return html

def build_attachments(jobs):
    attachments = []
    for i, job in enumerate(jobs):
        company_clean = job.get('company', 'Company').replace(' ', '_')

        resume_bytes = create_resume_doc(job, job.get('tailored_resume', ''))
        attachments.append((
            f"Job_{i+1}_{company_clean}_Resume.docx",
            resume_bytes
        ))

        cover_bytes = create_cover_letter_doc(job, job.get('cover_letter', ''))
        attachments.append((
            f"Job_{i+1}_{company_clean}_CoverLetter.docx",
            cover_bytes
        ))

    return attachments

def run_daily_job_search():
    print(f"\n[{datetime.now()}] Running scheduled job search...")
    criteria = """
LinkedIn ONLY. Canada ONLY.
Roles: Scrum Master, Project Manager
Location: Remote Canada or Hybrid Vancouver BC
Active postings ONLY — must be currently accepting applications (not closed, expired, or filled)
Salary requirement: minimum $75/hour or $180,000 annual — skip if salary is listed below this threshold; include if salary is not listed
Prioritize: newest postings first
"""
    try:
        result = find_and_tailor_jobs(criteria)
        jobs = result.get("jobs", [])
        summary = result.get("search_summary", "")

        if not jobs:
            print(f"No jobs found. Summary: {summary}")
            try:
                send_email(
                    f"Job Search Ran — No Results — {datetime.now().strftime('%b %d %Y')}",
                    f"""<html><body style='font-family:Arial;padding:20px;'>
<h2 style='color:#3b5bdb;'>Job Search Ran — No Results</h2>
<p>The search ran but found no matching Scrum Master or Project Manager roles on LinkedIn Canada today.</p>
<p style='color:#666;'>{summary}</p>
<p>Will try again at the next scheduled run.</p>
</body></html>""",
                    []
                )
            except Exception as mail_err:
                print(f"Could not send no-results email: {mail_err}")
            return

        html = build_email_html(jobs, summary)
        attachments = build_attachments(jobs)
        subject = f"LinkedIn Canada Jobs — {len(jobs)} matches — Newest first — {datetime.now().strftime('%b %d %Y')}"
        send_email(subject, html, attachments)
        print(f"Done. {len(jobs)} jobs found. Email sent with .docx attachments.")

    except Exception as e:
        error_msg = str(e)
        print(f"Scheduled job search failed: {error_msg}")
        if "429" in error_msg or "rate_limit" in error_msg:
            print("Rate limit — will retry at next scheduled run")

scheduler = BackgroundScheduler()

if SCHEDULER_ENABLED:
    scheduler.add_job(
        run_daily_job_search,
        trigger="cron",
        hour=8,
        minute=0,
        id="morning_search"
    )
    scheduler.add_job(
        run_daily_job_search,
        trigger="cron",
        hour=13,
        minute=0,
        id="afternoon_search"
    )
    print("Scheduler active — 8AM and 1PM daily")
else:
    print("Scheduler OFF — set SCHEDULER_ENABLED=true in .env to enable")

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/analyze", methods=["POST"])
def analyze():
    job_input = request.json.get("job_input", "").strip()
    if not job_input:
        return jsonify({"error": "Please paste a job description or URL."})
    try:
        return jsonify(analyze_job(job_input))
    except Exception as e:
        return jsonify({"error": str(e)})

@app.route("/find-jobs", methods=["POST"])
def find_jobs():
    criteria = request.json.get("criteria", "").strip()
    if not criteria:
        return jsonify({"error": "Please describe what jobs you are looking for."})
    try:
        return jsonify(find_and_tailor_jobs(criteria))
    except Exception as e:
        return jsonify({"error": str(e)})

@app.route("/trigger-now", methods=["POST"])
def trigger_now():
    thread = threading.Thread(target=run_daily_job_search)
    thread.daemon = True
    thread.start()
    return jsonify({"message": "Running in background — check your email in 2-3 minutes"})

@app.route("/scheduler-status")
def scheduler_status():
    return jsonify({
        "enabled": SCHEDULER_ENABLED,
        "message": "Scheduler ON — 8AM and 1PM daily" if SCHEDULER_ENABLED else "Scheduler OFF"
    })

if __name__ == "__main__":
    scheduler.start()
    print("Visit http://127.0.0.1:5000")
    try:
        app.run(debug=False)
    except (KeyboardInterrupt, SystemExit):
        scheduler.shutdown()